# pages/1_History_Dashboard.py
import sys
import os
# 상위 폴더(루트)의 auth 모듈을 import 할 수 있도록 경로 추가
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

import streamlit as st
import sqlite3
import pandas as pd
import json
from docx import Document
import io

from auth import require_login, render_sidebar_user_panel
from logging_setup import get_logger
from ui_theme import inject_global_css, page_header, status_badge

_log = get_logger("HISTORY")

st.set_page_config(page_title="생성 기록 대시보드", layout="wide")
inject_global_css()
# --- 인증 가드 ---
_user = require_login()
if st.session_state.get("_history_logged_user") != _user["email"]:
    _log.info(f"History 페이지 진입: {_user['email']}")
    st.session_state._history_logged_user = _user["email"]
render_sidebar_user_panel()

DB_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "history.db")

def load_projects_from_db(search_query=""):
    """데이터베이스에서 모든 프로젝트 목록을 검색어에 따라 불러옵니다."""
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    if search_query:
        c.execute("SELECT id, topic, timestamp FROM projects WHERE topic LIKE ? ORDER BY timestamp DESC", ('%' + search_query + '%',))
    else:
        c.execute("SELECT id, topic, timestamp FROM projects ORDER BY timestamp DESC")
    projects = [dict(row) for row in c.fetchall()]
    conn.close()
    return projects

def delete_project_from_db(project_id):
    """특정 프로젝트와 관련된 모든 데이터를 DB에서 삭제합니다."""
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("DELETE FROM projects WHERE id = ?", (project_id,))
    c.execute("DELETE FROM project_stages WHERE project_id = ?", (project_id,))
    conn.commit()
    conn.close()


def load_stages_from_db(project_id):
    """특정 프로젝트의 모든 단계별 결과를 불러옵니다."""
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT stage_name, content, llm_type FROM project_stages WHERE project_id = ? ORDER BY id ASC", (project_id,))
    stages = [dict(row) for row in c.fetchall()]
    conn.close()
    return stages

def create_docx_from_db(content):
    """텍스트 콘텐츠로 DOCX 파일을 생성합니다."""
    document = Document()
    for line in content.split('\n'):
        if line.startswith('### '):
            document.add_heading(line.lstrip('# ').strip(), level=3)
        elif line.startswith('## '):
            document.add_heading(line.lstrip('# ').strip(), level=2)
        elif line.startswith('# '):
            document.add_heading(line.lstrip('# ').strip(), level=1)
        elif line.strip():
            document.add_paragraph(line.strip())
    bio = io.BytesIO()
    document.save(bio)
    return bio.getvalue()

page_header(
    title="🗂️ 생성 기록 대시보드",
    subtitle="과거에 생성한 모든 제안서 프로젝트를 확인하고 관리합니다.",
)

search_query = st.text_input("주제 검색:", placeholder="검색할 제안서 주제를 입력하세요...")

projects = load_projects_from_db(search_query)

if not projects:
    st.info("아직 생성된 프로젝트 기록이 없거나, 검색 결과가 없습니다.")
else:
    st.caption(f"총 **{len(projects)}개** 프로젝트")

    # 게시판 헤더
    h1, h2, h3, h4 = st.columns([0.7, 4, 2, 2.4])
    with h1: st.markdown('<div class="appx-th">ID</div>', unsafe_allow_html=True)
    with h2: st.markdown('<div class="appx-th">주제</div>', unsafe_allow_html=True)
    with h3: st.markdown('<div class="appx-th">생성일</div>', unsafe_allow_html=True)
    with h4: st.markdown('<div class="appx-th appx-th-center">작업</div>', unsafe_allow_html=True)

    for project in projects:
        topic_html = (
            f'<span>{project["topic"]}</span>' if project["topic"]
            else '<span class="appx-untitled">(제목 미확정)</span>'
        )
        c1, c2, c3, c4 = st.columns([0.7, 4, 2, 2.4])
        with c1:
            st.markdown(f'<div class="appx-td appx-td-id">#{project["id"]}</div>', unsafe_allow_html=True)
        with c2:
            st.markdown(f'<div class="appx-td appx-td-topic">{topic_html}</div>', unsafe_allow_html=True)
        with c3:
            st.markdown(f'<div class="appx-td appx-td-date">{project["timestamp"]}</div>', unsafe_allow_html=True)
        with c4:
            bc1, bc2 = st.columns(2)
            with bc1:
                is_open = st.session_state.get("selected_project_id_for_view") == project['id']
                if st.button("닫기" if is_open else "상세", key=f"view_{project['id']}", use_container_width=True):
                    if is_open:
                        del st.session_state.selected_project_id_for_view
                    else:
                        st.session_state.selected_project_id_for_view = project['id']
                    st.rerun()
            with bc2:
                if st.button("삭제", key=f"delete_{project['id']}", use_container_width=True, type="secondary"):
                    delete_project_from_db(project['id'])
                    st.success(f"프로젝트 #{project['id']}가 삭제되었습니다.")
                    st.rerun()
        st.markdown('<div style="border-bottom:1px solid var(--appx-border);margin:0;"></div>', unsafe_allow_html=True)

        
        if st.session_state.get("selected_project_id_for_view") == project['id']:
            with st.expander(f"프로젝트 #{project['id']} 상세 내용", expanded=True):
                stages = load_stages_from_db(project['id'])
                if not stages:
                    st.info("이 프로젝트의 단계별 기록이 없습니다.")
                else:
                    stage_dict = {s['stage_name']: s for s in stages}

                    # ─────────────────────────────────────────────
                    # ① 다운로드 패널 (포인트 컬러 강조)
                    # ─────────────────────────────────────────────
                    draft = stage_dict.get("3단계: 본문 생성")
                    final = stage_dict.get("4단계: 최종본")
                    citations = stage_dict.get("3단계: 출처 목록")
                    date_prefix = project['timestamp'].split(' ')[0]
                    safe_topic = (project['topic'] or 'untitled').strip()

                    st.markdown(
                        '<div style="background:var(--appx-primary-soft);border-left:4px solid var(--appx-primary);'
                        'padding:10px 14px;border-radius:6px;margin-bottom:10px;">'
                        '<div style="font-weight:700;font-size:0.9rem;color:var(--appx-primary-hover);">📦 산출물 다운로드</div>'
                        '<div style="font-size:0.78rem;color:var(--appx-text-muted);margin-top:2px;">'
                        '생성된 단계별 파일을 즉시 받을 수 있습니다.</div>'
                        '</div>',
                        unsafe_allow_html=True,
                    )

                    d1, d2, d3, d4, d5 = st.columns(5)
                    with d1:
                        if draft:
                            st.download_button(
                                "📄 초안 DOCX",
                                data=create_docx_from_db(draft['content']),
                                file_name=f"[{date_prefix}] {safe_topic}_초안.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"dl_draft_docx_{project['id']}",
                                use_container_width=True,
                                type="primary",
                            )
                        else:
                            st.button("📄 초안 DOCX", key=f"dl_draft_docx_dis_{project['id']}", use_container_width=True, disabled=True)
                    with d2:
                        if draft:
                            st.download_button(
                                "📃 초안 TXT",
                                data=draft['content'].encode('utf-8'),
                                file_name=f"[{date_prefix}] {safe_topic}_초안.txt",
                                mime="text/plain",
                                key=f"dl_draft_txt_{project['id']}",
                                use_container_width=True,
                            )
                        else:
                            st.button("📃 초안 TXT", key=f"dl_draft_txt_dis_{project['id']}", use_container_width=True, disabled=True)
                    with d3:
                        if final:
                            st.download_button(
                                "✨ 최종본 DOCX",
                                data=create_docx_from_db(final['content']),
                                file_name=f"[{date_prefix}] {safe_topic}_최종본.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"dl_final_docx_{project['id']}",
                                use_container_width=True,
                                type="primary",
                            )
                        else:
                            st.button("✨ 최종본 DOCX", key=f"dl_final_docx_dis_{project['id']}", use_container_width=True, disabled=True)
                    with d4:
                        if final:
                            st.download_button(
                                "📋 최종본 TXT",
                                data=final['content'].encode('utf-8'),
                                file_name=f"[{date_prefix}] {safe_topic}_최종본.txt",
                                mime="text/plain",
                                key=f"dl_final_txt_{project['id']}",
                                use_container_width=True,
                            )
                        else:
                            st.button("📋 최종본 TXT", key=f"dl_final_txt_dis_{project['id']}", use_container_width=True, disabled=True)
                    with d5:
                        if citations:
                            st.download_button(
                                "🔗 출처 TXT",
                                data=citations['content'].encode('utf-8'),
                                file_name=f"[{date_prefix}] {safe_topic}_citations.txt",
                                mime="text/plain",
                                key=f"dl_cite_txt_{project['id']}",
                                use_container_width=True,
                            )
                        else:
                            st.button("🔗 출처 TXT", key=f"dl_cite_txt_dis_{project['id']}", use_container_width=True, disabled=True)

                    # ─────────────────────────────────────────────
                    # ② 후속 작업 네비게이션
                    # ─────────────────────────────────────────────
                    n1, n2, _ = st.columns([1, 1, 2])
                    with n1:
                        if st.button("🧐 품질 검증으로 이동", key=f"nav_review_{project['id']}", use_container_width=True, disabled=not draft):
                            st.session_state.selected_project_id = project['id']
                            st.session_state.active_tab = "4단계: 최종 품질 검증"
                            st.switch_page("0_Proposal_Generator.py")
                    with n2:
                        if st.button("📝 PPT 전환으로 이동", key=f"nav_ppt_{project['id']}", use_container_width=True, disabled=not (final or draft)):
                            st.session_state.selected_project_id = project['id']
                            st.session_state.active_tab = "5단계: PPT 전환"
                            st.switch_page("0_Proposal_Generator.py")

                    st.markdown('<div style="height:6px;"></div>', unsafe_allow_html=True)

                    # ─────────────────────────────────────────────
                    # ③ 단계별 결과 미리보기 (탭)
                    # ─────────────────────────────────────────────
                    tab_names = [s.split(': ')[1] for s in stage_dict.keys() if ':' in s]
                    stage_tabs = st.tabs(tab_names)

                    tab_idx = 0
                    for stage_name, stage in stage_dict.items():
                        if ':' not in stage_name: continue
                        with stage_tabs[tab_idx]:
                            st.caption(f"📌 {stage['stage_name']}" + (f"  ·  LLM: {stage['llm_type']}" if stage['llm_type'] else ""))
                            content = stage['content']
                            try:
                                parsed_json = json.loads(content)
                                st.json(parsed_json)
                            except (json.JSONDecodeError, TypeError):
                                st.text_area("내용", content, height=280, key=f"stage_{project['id']}_{tab_idx}", label_visibility="collapsed")
                        tab_idx += 1
